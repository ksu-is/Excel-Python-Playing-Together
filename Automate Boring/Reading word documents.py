#reading word documents

import docx

d = docx.Document('demo.docx')

#print(d.paragraphs)

print(d.paragraphs[0].text)#displays the text of the word document
print(d.paragraphs[1].text)
print(d.paragraphs[2].text)

p= d.paragraphs[1]
p.runs # numbers of "runs" of the different styles of text(regular, bold, Italic etc)
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)
p.runs[3].italic #in the shell this would show true as it is italic
p.runs[3].underline= True #this adds underline to the run when turn to True
p.runs[3].text=('italic and underline') #changes the text in the run
d.save('demo2.docx')  #saves the document

p.style #displays the style of the title, nowmal, heading 1, heading 2 etc.
p.style=('Title')#changes the style to title font
d.save('demo3.docx')

d=docx.Document()
d.add_paragraph('Hello, This is a paragraph')
d.add_paragraph('This is another paragraph I have added')
d.save('demo4.docx')

p=d.paragraphs[0]#how to add new paragraphs to document by adding runs
p.add_run('This is a new Run')
p.runs#displays the number of runs
p.runs[1].bold = True
d.save('demo5.docx')

#to edit an existing document you must create a new document 
# because you can only edit at the ends of runs

docx.Document()#add runs, paragrapghs etc to this doc from orignal
